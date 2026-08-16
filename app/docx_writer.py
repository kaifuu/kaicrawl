# -*- coding: utf-8 -*-
"""按《爬虫数据来源.xlsx》格式规范生成 WORD 文件。

规范：
  标题：加粗 / 宋体 / 五号(10.5pt) / 居中 / 单倍行距
  正文：首行空两格 / 宋体 / 五号 / 左对齐 / 单倍行距
  作者：单位动态 -> 「各单位」；其他用原作者（文章页「来源：xxx」抽取值亦计入作者）
  图片：png/jpg，按原文位置插入，并单独保存到当天目录的 images 子目录
"""
import base64
import os
from concurrent.futures import ThreadPoolExecutor

from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

from config import (OUTPUT_DIR, FONT_NAME, FONT_SIZE_PT, IMAGE_WORKERS,
                    IMAGE_WIDTH_RATIO, IMAGE_HEIGHT_RATIO)
from .utils import safe_filename, unique_path, ParserError
from .utils import http_get

IMG_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}

# 图片文件头魔数，用于校验下载内容确实是图片（防止把 404/HTML 错误页误存为图片）
_IMG_MAGICS = (b"\xff\xd8\xff", b"\x89PNG\r\n\x1a\n", b"GIF87a", b"GIF89a", b"BM")


def _is_image_bytes(data, ctype=""):
    """下载到的字节是否真的是图片：优先校验魔数，Content-Type 作兜底。"""
    if not data:
        return False
    if data.startswith(_IMG_MAGICS):
        return True
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return True
    return "image/" in (ctype or "").lower() and len(data) > 64


def _set_run_font(run, size_pt=FONT_SIZE_PT, bold=False):
    """设置 run 字体为宋体（含中文 eastAsia）、字号、加粗。"""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)


def _set_single_spacing(paragraph):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _set_first_line_indent_2chars(paragraph):
    """首行缩进 2 个中文字符（OOXML firstLineChars=200）。"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), "200")
    # 回退值（twips），2 字符 × 五号(10.5pt) × 20 ≈ 420
    ind.set(qn("w:firstLine"), str(int(2 * FONT_SIZE_PT * 20)))


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_single_spacing(p)
    run = p.add_run(text)
    _set_run_font(run, bold=True)
    return p


def _add_meta(doc, text):
    """作者/日期行：右对齐，宋体五号。"""
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_single_spacing(p)
    run = p.add_run(text)
    _set_run_font(run)
    return p


def _add_body_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_single_spacing(p)
    _set_first_line_indent_2chars(p)
    run = p.add_run(text)
    _set_run_font(run)
    return p


def _fit_image_inches(path, max_w, max_h):
    """按图片原生像素/DPI 求英寸尺寸，等比缩到 (max_w, max_h) 内，只缩不放。

    DPI 异常（如 1 或 600+，常见于某些截图工具）时按 96 处理。
    """
    with Image.open(path) as im:
        px_w, px_h = im.size
        try:
            dpi_x, dpi_y = im.info.get("dpi", (96, 96))
        except Exception:
            dpi_x = dpi_y = 96
        if not 24 <= dpi_x <= 600:
            dpi_x = 96
        if not 24 <= dpi_y <= 600:
            dpi_y = 96
    w, h = px_w / dpi_x, px_h / dpi_y
    if w <= 0 or h <= 0:
        return max_w, max_w * 0.6
    scale = min(max_w / w, max_h / h, 1.0)
    return w * scale, h * scale


def _add_image(doc, local_path):
    """居中插图：尺寸按本页可用区域的比例约束（宽 65% / 高 55%），不拉满整页。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_single_spacing(p)
    run = p.add_run()
    try:
        sec = doc.sections[0]
        # Length 相减返回裸 EMU int，需包回 Emu 再取英寸
        avail_w = Emu(sec.page_width - sec.left_margin - sec.right_margin).inches
        avail_h = Emu(sec.page_height - sec.top_margin - sec.bottom_margin).inches
        w, h = _fit_image_inches(local_path,
                                 avail_w * IMAGE_WIDTH_RATIO,
                                 avail_h * IMAGE_HEIGHT_RATIO)
        run.add_picture(local_path, width=Inches(w), height=Inches(h))
    except Exception:
        run.add_text("[图片插入失败]")
    return p


def _decode_data_url(url):
    """解码 data:image/png;base64,xxx 形式的内嵌图片，返回 (bytes, mime)。"""
    header, _, b64 = url.partition(",")
    mime = header[5:].split(";", 1)[0].strip().lower() or "image/png"
    data = base64.b64decode(b64 + "=" * (-len(b64) % 4))
    return data, mime


def _download_image(url, images_dir, seq, name_prefix=""):
    """下载单张图片到 images_dir，返回本地路径或 None。

    文件名：<name_prefix>_<seq>.<ext>。name_prefix 取自文章 docx 文件名，使每篇文章
    的图片在共享的 images/ 目录里互不撞名（旧实现用 001/002 会被同分类同日的其它文章覆盖）。
    支持三种 src：http(s) 远程图片、data:image/ 内嵌图片（手动导入粘贴的截图）。
    下载后校验字节确为图片（魔数），非图片（如 404/HTML 错误页）返回 None。
    """
    try:
        if url.startswith("data:"):
            data, ctype = _decode_data_url(url)
        else:
            resp = http_get(url)
            if resp.status_code >= 400:
                return None
            data = resp.content
            ctype = resp.headers.get("Content-Type", "")
        if not _is_image_bytes(data, ctype):
            return None
        # 由 Content-Type 或 URL 推断扩展名
        ctype_l = ctype.lower()
        ext = ".jpg"
        if "png" in ctype_l:
            ext = ".png"
        elif "jpeg" in ctype_l or "jpg" in ctype_l:
            ext = ".jpg"
        elif "webp" in ctype_l:
            ext = ".webp"
        elif "gif" in ctype_l:
            ext = ".gif"
        else:
            low = url.lower().split("?", 1)[0]
            for e in IMG_EXTS:
                if low.endswith(e):
                    ext = e
                    break
        os.makedirs(images_dir, exist_ok=True)
        prefix = safe_filename(name_prefix) if name_prefix else f"{seq:03d}"
        path = os.path.join(images_dir, f"{prefix}_{seq}{ext}")
        with open(path, "wb") as f:
            f.write(data)
        return path
    except Exception:
        return None


def _resolve_author(source, detail):
    """作者规则：单位动态 -> 各单位；其他 -> 原作者（含文章页「来源：xxx」抽取值）。"""
    if source.author_policy:
        return source.author_policy
    return (detail.get("author") or "").strip()


def generate(source, detail, date_str):
    """生成单篇文章的 WORD 文件。

    返回 (docx_abs_path, images_abs_dir)。
    落盘路径：OUTPUT_DIR/<分类>/<date>/<标题>.docx，图片存于同目录 images/。
    """
    category_dir = safe_filename(source.category) or "未分类"
    out_dir = os.path.join(OUTPUT_DIR, category_dir, date_str)
    images_dir = os.path.join(out_dir, "images")
    os.makedirs(out_dir, exist_ok=True)

    title = (detail.get("title") or "无标题").strip()
    base_name = safe_filename(title)
    # 先定下 docx 文件名（含可能的 _N 去重后缀）；图片以其为前缀命名，避免共享 images/ 目录撞名
    docx_path = unique_path(out_dir, base_name, ".docx")
    img_prefix = os.path.splitext(os.path.basename(docx_path))[0]

    doc = Document()
    # 默认正文字体也设为宋体五号，兜底
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_PT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    _add_title(doc, title)

    author = _resolve_author(source, detail)
    meta_parts = []
    if author:
        meta_parts.append(f"作者：{author}")
    if detail.get("publish_date"):
        meta_parts.append(detail["publish_date"])
    _add_meta(doc, "  ".join(meta_parts))

    # 第一遍：按块顺序给图片编号，并发预下载全部图片（_download_image 无状态、
    # 文件名含 img_prefix 互不冲突；序号/失败占号规则与旧的串行实现一致）
    img_jobs = []   # [(seq, src), ...]，顺序即块顺序
    for block in detail.get("blocks", []):
        if block.get("type") == "image" and (block.get("src") or ""):
            img_jobs.append((len(img_jobs) + 1, block["src"]))
    preloaded = {}  # seq -> 本地路径或 None
    if img_jobs:
        with ThreadPoolExecutor(max_workers=IMAGE_WORKERS, thread_name_prefix="img-dl") as ipool:
            futs = [ipool.submit(_download_image, src, images_dir, seq, name_prefix=img_prefix)
                    for seq, src in img_jobs]
            preloaded = {seq: fut.result() for (seq, _), fut in zip(img_jobs, futs)}

    # 第二遍：按原文顺序写入正文块（图片直接取预下载结果）
    img_seq = 0
    for block in detail.get("blocks", []):
        if block["type"] == "text":
            data = (block.get("data") or "").strip()
            if data:
                _add_body_text(doc, data)
        elif block["type"] == "image":
            src = block.get("src") or ""
            if not src:
                continue
            img_seq += 1
            local = preloaded.get(img_seq)
            if local:
                _add_image(doc, local)
            else:
                _add_body_text(doc, "[图片加载失败]")

    doc.save(docx_path)
    return docx_path, images_dir
